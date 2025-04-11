import cohere
from docx import Document

# Set up Cohere API client
cohere_api_key = "WQr8zIfWIlVTjOz5yZNMcum8XTuH1ERs62OiZDkz"
co = cohere.Client(cohere_api_key)

# Get title and synopsis from user input
title = "How to become an enduring conglomerate ASAP"
synopsis = "This nonfiction work would dissect what makes companies not only great but also capable of scaling rapidly in today’s fast-paced market. It combines deep, data-backed research with actionable frameworks and strategies that entrepreneurs and business leaders can implement immediately."

# Step 1: Generate Outline
def generate_outline(title, synopsis):
    prompt = f"""
    You are a team of writers led by Jim Collins and Alex Hormozi (imagine, not really)
    Create an outline for a novel titled '{title}', which is about {synopsis}. 
    Provide 15 cohesive chapter titles for this novel in quotation marks with 101st being the epilogue, followed by a detailed synopsis for each chapter in a structured format.
    Print only content
    Example:
    "Chapter 1: (chapter name)" - chapter synopsis...
    "Chapter 2: (chapter name)" - chapter synopsis...
    and onwards...
    Group 10 chapters as each part, each milestone in the story.
    Print content only"""import argparse
import pandas as pd
import os
from tabulate import tabulate

def load_data(file_path):
    """Load business process data from a CSV file."""
    try:
        df = pd.read_csv(file_path)
        print("Data loaded successfully.")
        return df
    except Exception as e:
        print(f"Error reading CSV file: {e}")
        exit(1)

def validate_data(df):
    """Check if required columns are present."""
    expected_columns = {'Process_Name', 'Task_Name', 'Department', 'Time_Taken', 'Frequency'}
    missing_columns = expected_columns - set(df.columns)
    if missing_columns:
        print(f"Missing required columns: {missing_columns}")
        exit(1)
    return True

def analyze_synergy(df):
    """
    Identify potential synergy points by grouping by Task_Name and counting unique departments.
    A task performed by more than one department indicates potential synergies.
    """
    synergy_df = df.groupby('Task_Name')['Department'].nunique().reset_index()
    synergy_df = synergy_df.rename(columns={'Department': 'Unique_Departments'})
    # Filter for tasks where more than one department is involved
    synergy_df = synergy_df[synergy_df['Unique_Departments'] > 1].sort_values(by='Unique_Departments', ascending=False)
    return synergy_df

def additional_analysis(df):
    """
    Compute additional metrics for each task:
      - Average time taken per task
      - Total time spent on the task
      - Average frequency and total frequency.
    """
    agg_df = df.groupby('Task_Name').agg({
        'Time_Taken': ['mean', 'sum'],
        'Frequency': ['mean', 'sum']
    })
    # Flatten the MultiIndex columns
    agg_df.columns = ['Time_Taken_Mean', 'Time_Taken_Sum', 'Frequency_Mean', 'Frequency_Sum']
    agg_df = agg_df.reset_index()
    return agg_df

def merge_analysis(synergy_df, agg_df):
    """Merge the synergy analysis with additional metrics."""
    merged_df = pd.merge(synergy_df, agg_df, on='Task_Name', how='left')
    return merged_df

def save_report(df, output_file):
    """Save the analysis report to a CSV file."""
    try:
        df.to_csv(output_file, index=False)
        print(f"Report saved to {output_file}")
    except Exception as e:
        print(f"Error saving report: {e}")

def print_report(df):
    """Print the DataFrame as a formatted table to the terminal."""
    print(tabulate(df, headers='keys', tablefmt='psql', showindex=False))

def main():
    parser = argparse.ArgumentParser(
        description="RPA Agent for Business Process Auditing & Synergy Analysis")
    parser.add_argument('--input', '-i', required=True,
                        help="Path to the input CSV file")
    parser.add_argument('--output', '-o', default="synergy_report.csv",
                        help="Path to save the output CSV report")
    parser.add_argument('--analysis', '-a', action='store_true',
                        help="Include additional analysis (time and frequency metrics)")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Input file {args.input} does not exist.")
        exit(1)
    
    # Step 1: Load and validate data
    df = load_data(args.input)
    validate_data(df)
    
    # Step 2: Analyze data to identify synergy points
    synergy_df = analyze_synergy(df)
    
    # Optional additional analysis
    if args.analysis:
        agg_df = additional_analysis(df)
        final_df = merge_analysis(synergy_df, agg_df)
    else:
        final_df = synergy_df
    
    # Step 3: Output results to terminal and CSV
    print("\nSynergy Analysis Report:")
    print_report(final_df)
    save_report(final_df, args.output)

if __name__ == "__main__":
    main()
    
    response = co.generate(
        model='command-r-plus',  # Use the latest available model
        prompt=prompt,
    )
    outline = response.generations[0].text.strip()
    return outline

outline = generate_outline(title, synopsis)
print("\nBook Outline Generated:")
print(outline)

# Step 2: Write the Book Piece by Piece
def write_chapter(chapter_title):
    prompt = f"""You are a team of writers led by Jim Collins and Alex Hormozi (imagine, not really). Write a detailed chapter titled '{chapter_title}'.
    Make every tough concept easy to understand with simple, real-life examples. The chapter should be long and detailed, ensuring a compelling narrative, and as long as can be possible. I mean make it inhumanly long.
    This book is for MBA grads and Aspiring Entrepreneurs. So, go over all sorts of startegi, and test them with everything you got to find the ideal mix of strategies to achieve the fastest sustainable growth model.
    Write only the chapter content, no need to label the chapters as Chapter 1: content or anything, that is taken care of. Also ensure that each delves into great detail, with case studies, best practices, and more.
    (Print content only)."""
    
    response = co.generate(
        model='command-r-plus',
        prompt=prompt,
    )
    chapter_content = response.generations[0].text.strip()
    return chapter_content

# Extract chapter titles correctly
import re
chapter_pattern = re.findall(r'"(.*?)" - (.*)', outline)

book_content = {}
for chapter_title, chapter_summary in chapter_pattern:
    print(f"\nWriting Chapter: {chapter_title}")
    chapter_content = write_chapter(chapter_title)
    book_content[chapter_title] = chapter_content

# Step 3: Save the Content into a Word File
def save_to_word(book_content, title):
    doc = Document()
    doc.add_heading(title, level=1)

    for chapter, content in book_content.items():
        doc.add_heading(f"{chapter}", level=2)
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save(f"{title}.docx")
    print(f"\nBook saved as {title}.docx")

save_to_word(book_content, title)